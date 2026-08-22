"""Build the Google Docs-ready Week 2 submission document."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(31, 41, 55)
MUTED = RGBColor(75, 85, 99)
ACCENT = RGBColor(15, 118, 110)
ACCENT_DARK = RGBColor(17, 94, 89)
LIGHT_TEAL = "E6F4F1"
LIGHT_BLUE = "EAF1F8"
LIGHT_GRAY = "F3F4F6"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    tr_pr.append(table_header)


def set_fixed_table_layout(table) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_inches: float) -> None:
    width_twips = int(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.widow_control = True

    title = document.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(26)
    title.font.bold = False
    title.font.color.rgb = INK
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)

    for style_name, size, before, after, color in (
        ("Heading 1", 20, 20, 6, INK),
        ("Heading 2", 16, 18, 6, INK),
        ("Heading 3", 14, 16, 4, RGBColor(67, 67, 67)),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.font.color.rgb = INK
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def add_label(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text.upper())
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = ACCENT


def add_body(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        paragraph.add_run(text[len(bold_lead) :])
    else:
        paragraph.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_callout(document: Document, title: str, body: str, fill: str = LIGHT_TEAL) -> None:
    table = document.add_table(rows=1, cols=1)
    set_fixed_table_layout(table)
    table.columns[0].width = Inches(6.5)
    set_cell_width(table.cell(0, 0), 6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 140, 160, 140, 160)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    title_run = paragraph.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = ACCENT_DARK
    body_paragraph = cell.add_paragraph(body)
    body_paragraph.paragraph_format.space_after = Pt(0)
    prevent_row_split(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    set_fixed_table_layout(table)
    for column, width in zip(table.columns, widths, strict=True):
        column.width = Inches(width)
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for index, (label, width) in enumerate(zip(headers, widths, strict=True)):
        cell = header.cells[index]
        set_cell_width(cell, width)
        set_cell_shading(cell, "0F766E")
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for index, (value, width) in enumerate(zip(values, widths, strict=True)):
            cell = row.cells[index]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            set_cell_shading(cell, "FFFFFF" if row_index % 2 == 0 else LIGHT_GRAY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            run.font.size = Pt(9.5)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_break(document: Document) -> None:
    document.add_page_break()


def build_document(output: Path) -> None:
    document = Document()
    configure_document(document)
    document.core_properties.title = "Regulation Navigator — Week 2 Project Submission"
    document.core_properties.subject = "RAG application project overview and evaluation"
    document.core_properties.author = "Regulation Navigator project team"

    add_label(document, "Week 2 · RAG application · August 2026")
    title = document.add_paragraph(style="Title")
    title.add_run("Medical Device Software\nRegulatory Navigator")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Project overview, system design, evaluation, and submission record")
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED

    add_callout(
        document,
        "One-line RAG application statement",
        "My RAG app helps medical-device regulatory, quality, and software professionals "
        "answer software-classification and applicable-requirement questions from 21 curated, "
        "English-language FDA, U.S. statutory, EU MDR, IMDRF, IEC, and ISO evidence cards in a "
        "Streamlit application with ≥95% claim-level faithfulness and p95 latency ≤5 seconds.",
    )

    document.add_heading("Executive summary", level=1)
    add_body(
        document,
        "The navigator screens one software function at a time and produces a preliminary, cited "
        "view of software type, U.S. FDA device status, Non-Device CDS eligibility, IMDRF SaMD "
        "category, IEC 62304 safety class, EU MDR Rule 11 class, and applicable lifecycle "
        "frameworks. Typed LangGraph nodes keep deterministic regulatory logic separate from "
        "retrieval, citation validation, claim grounding, refusal behavior, and presentation.",
    )
    add_body(
        document,
        "The intended users are medical-device regulatory, quality, and software professionals "
        "performing early product-definition or change-impact triage. The tool is educational "
        "research support—not legal advice, an agency determination, or a conformity assessment.",
    )
    add_table(
        document,
        ["Current development benchmark", "Result"],
        [
            ["Claim-level faithfulness", "100% on 20 end-to-end cases"],
            ["Retrieval quality", "100% Recall@5 · 0.933 MRR · 0.951 nDCG@5"],
            ["Refusal accuracy", "100% on supported and unsupported queries"],
            ["Latency", "0.004s p50 · 0.005s p95 · target ≤5s"],
        ],
        [4.25, 2.25],
    )

    add_page_break(document)
    document.add_heading("Required RAG framework", level=1)
    framework_rows = [
        [
            "Use case",
            "Answer preliminary software qualification, classification, and applicable-framework questions for one medical-software function, with cited claims, caveats, and missing-fact questions in Streamlit.",
        ],
        [
            "Corpus",
            "21 curated English JSONL evidence cards spanning FDA, U.S. statutes, EU MDR, IMDRF, IEC 62304, and ISO 14971; authorized PDF, Markdown, and text can also be ingested. Issuers own official truth; the maintainer owns snapshots and review metadata.",
        ],
        [
            "Ingestion + cleaning",
            "Extract PDF pages or Markdown heading paths, normalize whitespace, split at semantic boundaries, attach lifecycle metadata, reject empty/incomplete records, and deduplicate stable chunk IDs.",
        ],
        [
            "Ingestion + freshness",
            "Review every source family monthly under a 30-day SLA, with a one-business-day target for known material updates. A manifest audit blocks answer release when stale or incomplete.",
        ],
        [
            "Chunking + embedding",
            "Use hierarchy-aware 1,600-character chunks with 200-character overlap and a reproducible 384-dimensional signed hashing embedding. The size preserves coherent provisions; the offline model supports a transparent classroom baseline.",
        ],
        [
            "Retrieve",
            "Use local Chroma plus BM25, 56/44 reciprocal-rank fusion, exact-identifier boosts, top-k 8, and metadata-constrained supplements for each decision-critical source family.",
        ],
    ]
    add_table(document, ["Field", "Implementation"], framework_rows, [1.45, 5.05])

    document.add_heading("Source lifecycle and licensing", level=2)
    add_body(
        document,
        "Each evidence card records a stable ID, source family, jurisdiction, authority, document "
        "type, legal status, issue date, section, official URL, supersession flag, topics, and a "
        "concise evidence summary. Binding law, final guidance, draft guidance, harmonized "
        "frameworks, and licensed standards remain visibly distinct.",
    )
    add_callout(
        document,
        "Licensed standards boundary",
        "IEC and ISO clause text is intentionally not reproduced. The app exposes scope-level "
        "catalog evidence and directs users to licensed copies for clause-level conclusions.",
        LIGHT_BLUE,
    )

    document.add_heading("System architecture", level=1)
    add_body(
        document,
        "The workflow is a typed state machine. Classification is deterministic; LangChain handles "
        "documents, embeddings, Chroma, and retrieval, while LangGraph makes routing and evidence "
        "gates explicit and testable.",
    )
    architecture_rows = [
        [
            "1",
            "Input + facts",
            "Validate the Pydantic schema; extract only explicit or high-confidence facts.",
        ],
        ["2", "Scope + route", "Classify software scope and select CDS or general U.S. analysis."],
        [
            "3",
            "Framework screens",
            "Apply U.S., IMDRF N10/N12, IEC 62304, and EU MDR Rule 11 rules.",
        ],
        ["4", "Uncertainty", "Preserve unknown facts as focused follow-up questions."],
        [
            "5",
            "Hybrid retrieval",
            "Fuse BM25 and Chroma ranks; supplement controlling source families.",
        ],
        [
            "6",
            "Evidence validation",
            "Reject superseded or metadata-incomplete citations and audit freshness.",
        ],
        [
            "7",
            "Claim grounding",
            "Bind every displayed conclusion to required source keys and citation IDs.",
        ],
        [
            "8",
            "Release + render",
            "Answer, answer with caveats, or refuse; show links and quality metrics.",
        ],
    ]
    add_table(
        document, ["Step", "Node group", "Responsibility"], architecture_rows, [0.55, 1.65, 4.3]
    )

    document.add_heading("Refusal and confidence policy", level=2)
    add_bullets(
        document,
        [
            "Refuse unrelated or explicitly unsupported topics such as recipes, HIPAA, billing, or reimbursement.",
            "Refuse when atomic-claim source coverage is below 80% or the freshness audit fails.",
            "Return “answer with caveats” when evidence is sufficient but material product facts remain unresolved.",
            "Use ≥95% mean claim-level faithfulness as the project release target; keep the operational per-answer gate at 80% to block severely incomplete multi-framework results.",
        ],
    )

    document.add_heading("Prompts and agent instructions", level=1)
    document.add_heading("Fact extraction system instruction", level=2)
    add_callout(
        document,
        "Constrained extraction",
        "Extract only facts about one medical-software function that the description actually "
        "supports. Use null or ‘unknown’ for missing facts; do not infer clinical severity, "
        "hardware independence, or whether an HCP can independently review a recommendation.",
        LIGHT_BLUE,
    )
    document.add_heading("Fact extraction user prompt", level=2)
    add_body(
        document,
        "Provide the software description and intended-use statement, then request the constrained "
        "AssessmentInput schema. Explicit user-entered form facts override extracted candidates.",
    )
    document.add_heading("Grounded assessment instruction", level=2)
    add_callout(
        document,
        "Evidence discipline",
        "Use only deterministic classifications and supplied evidence; label conclusions "
        "preliminary; cite with [chunk_id]; distinguish law, final guidance, draft guidance, and "
        "harmonized frameworks; never reconstruct licensed IEC/ISO text; preserve missing facts.",
    )
    document.add_heading("Implementation note", level=2)
    add_body(
        document,
        "The default application works without an API key. Optional OpenAI structured extraction "
        "turns natural language into candidate facts but cannot decide applicability and falls "
        "back to deterministic extraction if unavailable. Grounded prose is generated "
        "deterministically from atomic claims so citation validation does not require an LLM judge.",
    )

    document.add_heading("Safety contract", level=2)
    add_bullets(
        document,
        [
            "Every inline citation ID must exist in the validated evidence set.",
            "Every atomic claim declares the source families required to support it.",
            "Draft guidance is labeled draft and never represented as binding law.",
            "IMDRF categories are not presented as FDA or EU device classes.",
            "Missing product facts remain visible rather than being silently inferred.",
        ],
    )

    add_page_break(document)
    document.add_heading("Evaluation", level=1)
    add_body(
        document,
        "The checked-in benchmark includes 20 end-to-end classification cases and 15 raw "
        "retrieval/refusal cases. Metrics below come from the 2026-08-22 local run; they are "
        "development results, not independent regulatory validation.",
    )
    add_table(
        document,
        ["Metric", "Measured", "Target / interpretation"],
        [
            ["Software-category accuracy", "100%", "Expected label substring across 20 cases"],
            ["U.S.-status accuracy", "100%", "Expected status across 20 cases"],
            ["Required-source recall", "100%", "All case-required source families present"],
            [
                "Claim-level faithfulness",
                "100%",
                "Target ≥95%; required citations per atomic claim",
            ],
            ["Inline citation validity", "100%", "No invented citation IDs"],
            ["Retrieval Recall@5", "100%", "Relevant source family in top five"],
            ["MRR / nDCG@5", "0.933 / 0.951", "Rank sensitivity beyond simple recall"],
            ["Refusal accuracy", "100%", "Supported vs. unsupported query gate"],
            ["p50 / p95 latency", "0.004s / 0.005s", "p95 target ≤5s; cold load ≈0.257s"],
        ],
        [2.15, 1.25, 3.1],
    )

    document.add_heading("Failure analysis", level=2)
    add_bullets(
        document,
        [
            "No automated release-check failures occurred in the current development sets.",
            "Source-family supplements protect completeness but can hide ranking weakness; raw retrieval is therefore measured separately.",
            "Exact 520(o) and OTS queries ranked the correct source second, explaining MRR and nDCG below 1.0 despite full Recall@5.",
            "The small curated benchmark can overstate generalization; blind expert labeling and adversarial cases are still required.",
            "The offline hashing embedding is a transparent baseline, not a production semantic model.",
        ],
    )

    add_page_break(document)
    document.add_heading("Iterations tried", level=1)
    iteration_rows = [
        [
            "1 · Rule-only result",
            "Deterministic classifications plus an evidence list.",
            "Reproducible, but prose claims were not explicitly bound to citations and evaluation was too coarse.",
        ],
        [
            "2 · Semantic baseline",
            "Local Chroma with signed hashing embeddings.",
            "Offline semantic retrieval worked, but exact regulatory identifiers could be displaced.",
        ],
        [
            "3 · Hybrid retrieval",
            "BM25 + vector RRF, exact boosts, and source-family supplements.",
            "Improved direct retrieval and guaranteed controlling-source completeness while keeping raw ranking measurable.",
        ],
        [
            "4 · Release gates",
            "Atomic claims, citation validation, freshness audit, refusal policy, and latency metrics.",
            "Closed the faithfulness, I-don’t-know, freshness, latency, and failure-analysis criteria.",
        ],
    ]
    add_table(document, ["Iteration", "Change", "Observation"], iteration_rows, [1.5, 2.15, 2.85])

    document.add_heading("Learnings and observations", level=1)
    add_bullets(
        document,
        [
            "Regulatory RAG needs semantic matching and exact retrieval because identifiers and framework names often carry more meaning than nearby prose.",
            "Deterministic nodes expose ambiguity; an LLM is more useful for constrained extraction and presentation than legal applicability.",
            "Source recall is not sufficient: every material claim must be bound to evidence that supports that specific claim.",
            "Publication date and index-review date answer different lifecycle questions and must both be tracked.",
            "Perfect results on a small curated set require skepticism, independent review, and a much larger blind benchmark.",
        ],
    )

    add_page_break(document)
    document.add_heading("Reproduction and handoff", level=1)
    add_body(document, "Run the application and release checks from the project root:")
    code = document.add_table(rows=1, cols=1)
    set_fixed_table_layout(code)
    code.columns[0].width = Inches(6.5)
    set_cell_width(code.cell(0, 0), 6.5)
    cell = code.cell(0, 0)
    set_cell_shading(cell, "111827")
    set_cell_margins(cell, 140, 160, 140, 160)
    commands = (
        "python3.12 -m venv .venv\n"
        "source .venv/bin/activate\n"
        'python -m pip install -e ".[dev]"\n'
        "streamlit run app.py\n"
        "pytest\n"
        "python -m regulation_navigator.evaluate\n"
        "regnav audit-corpus"
    )
    run = cell.paragraphs[0].add_run(commands)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = WHITE
    prevent_row_split(code.rows[0])

    document.add_heading("Submission artifacts", level=2)
    add_bullets(
        document,
        [
            "Runnable Streamlit application and LangGraph workflow.",
            "21-record source-linked corpus, ingestion CLI, and freshness manifest.",
            "20 classification cases, 15 retrieval/refusal cases, and a generated evaluation report.",
            "This Google Docs-ready project document, demo script, and ≤5-minute MP4 walkthrough.",
            "Repository contents ready for GitHub; publishing requires an authenticated destination and visibility choice.",
        ],
    )

    document.add_heading("AI coding tools used", level=2)
    add_body(
        document,
        "Codex was used to scaffold and implement the Python project, inspect the handout, refine "
        "the LangGraph workflow, add automated evaluation, run tests and static checks, and "
        "generate submission artifacts. Human review remains required for regulatory "
        "interpretations, source licensing, benchmark labels, and final submission decisions.",
    )

    add_callout(
        document,
        "Required human review",
        "The navigator does not determine U.S. product code, FDA class, submission pathway, EU "
        "conformity route, clinical-evidence sufficiency, or final IEC 62304 class. Exact claims, "
        "architecture, hazards, risk controls, current source text, and qualified regulatory "
        "review remain necessary.",
        LIGHT_BLUE,
    )

    for section in document.sections:
        section.start_type = WD_SECTION.NEW_PAGE
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build_document(args.output)


if __name__ == "__main__":
    main()
